from docx import Document
from openai import OpenAI

from docx import Document

# 打开并读取.docx文件
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)




def translate_document(input_path, output_path):

    content = read_docx(input_path)
    print('-------------------')
    print(content)
    print(len(content))
    print('-------------------')

    API_BASE = "https://api.siliconflow.cn/v1"
    API_KEY = "sk-wlyagmcadnehxxofddguvwtfngphbigvbshaufpwsysiihrj"
    client = OpenAI(
            api_key=API_KEY,
            base_url=API_BASE
        )
    completion = client.chat.completions.create(
        model = "Qwen/Qwen2-7B-Instruct",
        messages=[{"role": "user", "content": content}]
        )
    response = completion.choices[0].message.content
    print(response)
   
    # 遍历每一个段落并翻译
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # 如果段落不是空的
            translated_text = translator.translate(paragraph.text, src='zh-cn', dest='en').text
            translated_doc.add_paragraph(translated_text)
        else:
            translated_doc.add_paragraph('')  # 保持段落间的空行

    # 将翻译后的内容保存到新文档中
    translated_doc.save(output_path)
    print(f"Translation completed and saved to {output_path}")

# 设置输入和输出文档路径
input_path = "F:/code/large_model_test/data/Chinese-choice.docx"
output_path = "F:/code/large_model_test/output_result/translated_choice_document.docx"

# 调用函数进行翻译
translate_document(input_path, output_path)
